from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from docx2pdf import convert
import os

from formatter import format_sections  # ✅ Required for auto conversion

def export_to_docx(title: str, content, save_path: str = "./exports", company_1="Company 1", company_2="Company 2") -> str:
    os.makedirs(save_path, exist_ok=True)
    doc = Document()

    # ─── Title ─────────────────────
    heading = doc.add_paragraph()
    heading_run = heading.add_run(title.upper())
    heading_run.bold = True
    heading_run.font.size = Pt(18)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ─── Intro Clause ──────────────
    today = datetime.today().strftime("%B %d, %Y")
    intro = doc.add_paragraph()
    intro.add_run(
        f'\nThis Non-Disclosure Agreement (this "Agreement") is made effective as of {today} '
        f'(the "Effective Date"), by and between {company_1}, of [Owner Address], and {company_2}, of [Recipient Address].\n\n'
    ).font.size = Pt(12)

    p = doc.add_paragraph()
    r = p.add_run(
        "The Owner has requested and the Recipient agrees that the Recipient will protect the confidential "
        "material and information which may be disclosed between the Owner and the Recipient. "
        "Therefore, the parties agree as follows:"
    )
    r.font.size = Pt(12)

    # ─── Normalize Content ─────────
    if isinstance(content, str):
        content = format_sections(content)

    # ─── Add Clauses ───────────────
    for clause in content:
        text = clause["text"]
        for line in text.split("\n"):
            line = line.strip()

            if not line:
                continue
            elif line.startswith(". ") or line.isupper():
                p = doc.add_paragraph()
                r = p.add_run(line.lstrip(". ").upper())
                r.bold = True
                r.font.size = Pt(12)
            elif len(line) > 3 and line[1:3] == ". " and line[0].isalpha():
                p = doc.add_paragraph()
                r = p.add_run(line)
                r.bold = True
                r.font.size = Pt(11)
                p.paragraph_format.left_indent = Inches(0.3)
            elif line.startswith("- ") or line.startswith("• "):
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.5)
                p.add_run(line[2:])
            else:
                p = doc.add_paragraph()
                p.add_run(line).font.size = Pt(11)

    # ─── Signature Block ───────────
    doc.add_paragraph("\nIN WITNESS WHEREOF, the parties have executed this Agreement as of the Effective Date.\n")
    doc.add_paragraph(f"{company_1.upper()}:")
    doc.add_paragraph("By: ____________________________    Date: ____________________\n")
    doc.add_paragraph(f"{company_2.upper()}:")
    doc.add_paragraph("By: ____________________________    Date: ____________________")

    # ─── Save DOCX ─────────────────
    filename = f"NDA_{datetime.today().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = os.path.join(save_path, filename)
    doc.save(file_path)
    return file_path


def export_to_pdf(title: str, content, save_path: str = "./exports", company_1="Company 1", company_2="Company 2") -> str:
    docx_path = export_to_docx(title, content, save_path, company_1, company_2)
    pdf_path = docx_path.replace(".docx", ".pdf")
    convert(docx_path, pdf_path)
    return pdf_path
